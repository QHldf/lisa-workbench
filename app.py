# -*- coding: utf-8 -*-
"""
WorkBuddy · 融资顾问 Web App
后端：Flask + JSON 数据源（单一真相源 = 磁盘 JSON，全部内容实时同步）
核心诉求：快速查客户 / 快速查银行 / 银行政策改了即同步
同步机制：@app.before_request 每次请求前从磁盘重载全部数据，零重启、零缓存陈旧
"""
import json, os, re
from datetime import datetime, date
from flask import Flask, jsonify, request, render_template, send_file
from docx import Document
from io import BytesIO

BASE = os.path.dirname(os.path.abspath(__file__))

def find_kb():
    """兼容开发环境(/workspace/knowledge-base)与部署环境(数据可能被打包进 lisa-app)"""
    candidates = [
        os.path.join(os.path.dirname(BASE), 'knowledge-base'),
        os.path.join(BASE, 'knowledge-base'),
        '/workspace/knowledge-base',
    ]
    for c in candidates:
        if os.path.isdir(c):
            return c
    return candidates[0]

KB = find_kb()
BANK_DIR = os.path.join(KB, 'bank-data')

app = Flask(__name__, template_folder='templates', static_folder='static')

# ---- 路径 ----
WB_PATH = os.path.join(BANK_DIR, 'workbench_data.json')
PRODUCTS_PATH = os.path.join(BANK_DIR, 'products.json')
RULES_PATH = os.path.join(BANK_DIR, 'rules.json')

# ---- 内存缓存（每次请求前由 load_all 刷新）----
WB = {'clients': [], 'insts': [], 'active': []}
PRODUCTS = []
RULES = {}
ACTIVE = []
BANKS = []
CLIENTS_BRIEF = []
ACTIVE_NAMES = set()

def load_json(path, default=None):
    try:
        with open(path, encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        return default if default is not None else {}

# ============ 模板文本自动解析（模板数据→银行判断标准）============
def parse_template(text):
    """从客户模板纯文本中正则提取结构化字段，返回dict。未提取到的字段不在dict中。"""
    if not text:
        return {}
    r = {}
    # 法人年龄: "49岁" / "（49岁）"
    m = re.search(r'[（(]?\s*(\d{2})\s*岁\s*[)）]?', text)
    if m:
        r['legal_age'] = int(m.group(1))
    # 成立年份: "2021年成立"
    m = re.search(r'(\d{4})年\d*月?\s*成立', text)
    if m:
        r['established_year'] = int(m.group(1))
    # 持股比例
    m = re.search(r'(?:持股|占股)\s*(\d+(?:\.\d+)?)\s*%', text)
    if m:
        r['shareholding_pct'] = float(m.group(1))
    # 纳税级别: "纳税A级" / "纳税信用A级"
    m = re.search(r'纳税(?:信用)?\s*([ABCMD])\s*级', text)
    if m:
        r['tax_grade'] = m.group(1)
    # 近3月查询
    m = re.search(r'近\s*3\s*个?月.*?查询\s*(\d+)\s*次', text)
    if m:
        r['query_3m'] = int(m.group(1))
    # 近6月查询
    m = re.search(r'近\s*6\s*个?月.*?查询\s*(\d+)\s*次', text)
    if m:
        r['query_6m'] = int(m.group(1))
    # 近12月查询
    m = re.search(r'近\s*(?:12\s*个?月|一年).*?查询\s*(\d+)\s*次', text)
    if m:
        r['query_12m'] = int(m.group(1))
    # 信用卡使用率
    m = re.search(r'使用率(?:约)?\s*([\d.]+)\s*%', text)
    if m:
        r['cc_usage'] = float(m.group(1))
    # 销贷比
    m = re.search(r'销贷比(?:约|略超)?\s*([\d.]+)\s*%', text)
    if m:
        r['sale_ratio'] = float(m.group(1))
    # 逾期判断
    overdue_section = ''
    for line in text.split('\n'):
        if '征信逾期' in line or ('逾期' in line and '查询' not in line):
            overdue_section = line
            break
    if any(kw in overdue_section for kw in ['无', '暂无', '待补充']):
        r['has_overdue'] = False
    elif any(kw in overdue_section for kw in ['有', '连', '累']):
        r['has_overdue'] = True
    # 连三累六
    r['overdue_cont'] = 3 if re.search(r'连\s*三|连\s*3', text) else 0
    r['overdue_cum'] = 6 if re.search(r'累\s*六|累\s*6', text) else 0
    # 企业总负债: "借贷余额共计50968.39万" 或 "合计约400万"
    m = re.search(r'借贷余额共计\s*([\d,.]+)\s*万', text)
    if m:
        r['debt_total'] = float(m.group(1).replace(',', ''))
    else:
        m = re.search(r'合计(?:约)?\s*([\d,.]+)\s*万', text)
        if m:
            r['debt_total'] = float(m.group(1).replace(',', ''))
    # 是否双签
    r['is_dual_sign'] = '双签' in text
    return r

# ============ 同步核心：每次请求前重载全部数据 ============
def load_all():
    """从磁盘重新加载全部数据 —— 工作台所有内容（客户/银行/产品/分层/获贷率/红线/在做进展）实时同步，无需重启"""
    global WB, PRODUCTS, RULES, ACTIVE, BANKS, CLIENTS_BRIEF, ACTIVE_NAMES
    WB = load_json(WB_PATH, {'clients': [], 'insts': [], 'active': []})
    PRODUCTS = load_json(PRODUCTS_PATH, [])
    RULES = load_json(RULES_PATH, {})
    # 在做客户进展优先读 workbench_data.json 的 active[]（数据化，可同步）
    ACTIVE = WB.get('active', [])
    BANKS = build_banks()
    CLIENTS_BRIEF = build_clients_brief()
    ACTIVE_NAMES = {a['name'] for a in ACTIVE}

@app.before_request
def _refresh():
    load_all()

# ============ 持久化 ============
def save_wb():
    with open(WB_PATH, 'w', encoding='utf-8') as f:
        json.dump(WB, f, ensure_ascii=False, indent=2)

def save_products():
    with open(PRODUCTS_PATH, 'w', encoding='utf-8') as f:
        json.dump(PRODUCTS, f, ensure_ascii=False, indent=2)

def save_rules():
    with open(RULES_PATH, 'w', encoding='utf-8') as f:
        json.dump(RULES, f, ensure_ascii=False, indent=2)

# ============ 构建缓存 ============
def build_banks():
    tiers = RULES.get('tiers', {})
    rates = RULES.get('rates', {})
    redline = set(RULES.get('redline', []))
    banks = []
    for b in WB.get('insts', []):
        name = b['name']
        prods = [p for p in PRODUCTS if p.get('bank') == name]
        banks.append({
            'name': name,
            'count': b.get('count', 0),
            'tier': tiers.get(name, '其他'),
            'rate': rates.get(name, ''),
            'red': name in redline,
            'product_count': len(prods),
            'products': prods,
        })
    return banks

def build_clients_brief():
    out = []
    for c in WB.get('clients', []):
        out.append({
            'name': c['name'],
            'tag': c.get('tag', ''),
            'income': c.get('income'),
            'tax': c.get('tax'),
            'biz_inst_count': c.get('biz_inst_count', 0),
            'bank_count': c.get('bank_count', 0),
            'all_institutions': c.get('all_institutions', []),
            'active': c['name'] in ACTIVE_NAMES,
            'created': c.get('created', ''),
            'status': c.get('status', '历史'),
            # 结构化风控字段（供匹配引擎一键带出）
            'legal_age': c.get('legal_age'),
            'tax_grade': c.get('tax_grade', ''),
            'debt_total': c.get('debt_total'),
            'query_3m': c.get('query_3m'),
            'query_6m': c.get('query_6m'),
            'query_12m': c.get('query_12m'),
            'cc_usage': c.get('cc_usage'),
            'sale_ratio': c.get('sale_ratio'),
            'has_overdue': c.get('has_overdue', False),
            'overdue_cont': c.get('overdue_cont', 0),
            'overdue_cum': c.get('overdue_cum', 0),
            'is_gaoxin': c.get('is_gaoxin', False),
            'has_tech': c.get('has_tech', False),
        })
    return out

# ============ 路由 ============
@app.route('/')
def index():
    stats = {
        'total_clients': len(WB['clients']),
        'total_banks': len(BANKS),
        'total_products': len(PRODUCTS),
        'active_count': len(ACTIVE),
        'urgent_count': sum(1 for a in ACTIVE if a['status'] == '紧急'),
        'overdue_count': sum(1 for a in ACTIVE if a.get('overdue')),
    }
    return render_template('index.html', stats=stats)

@app.route('/api/stats')
def api_stats():
    return jsonify({
        'total_clients': len(WB['clients']),
        'total_banks': len(BANKS),
        'total_products': len(PRODUCTS),
        'active_count': len(ACTIVE),
        'urgent_count': sum(1 for a in ACTIVE if a['status'] == '紧急'),
        'overdue_count': sum(1 for a in ACTIVE if a.get('overdue')),
        'income_coverage': sum(1 for c in WB['clients'] if c.get('income')),
    })

@app.route('/api/active')
def api_active():
    today = date.today().isoformat()
    order = {'紧急': 0, '推进中': 1, '待反馈': 2, '暂停': 3}
    def sort_key(a):
        d = a.get('deadline')
        if d:
            try:
                dl = (date.fromisoformat(d) - date.today()).days
            except Exception:
                dl = 999
        else:
            dl = 999
        return (dl, order.get(a['status'], 9))
    data = sorted(ACTIVE, key=sort_key)
    return jsonify({'today': today, 'items': data})

@app.route('/api/clients')
def api_clients():
    q = request.args.get('q', '').strip().lower()
    tag = request.args.get('tag', '')
    bank = request.args.get('bank', '')
    active_only = request.args.get('active') == '1'
    result = CLIENTS_BRIEF
    if active_only:
        result = [c for c in result if c['active']]
    if tag:
        result = [c for c in result if tag in (c['tag'] or '')]
    if bank:
        result = [c for c in result if bank in (c['all_institutions'] or [])]
    if q:
        result = [c for c in result if q in (c['name'] + ' ' + (c['tag'] or '') + ' ' + ' '.join(c['all_institutions'] or [])).lower()]
    return jsonify(result)

@app.route('/api/client/<path:name>')
def api_client_detail(name):
    for c in WB['clients']:
        if c['name'] == name:
            progress = next((a for a in ACTIVE if a['name'] == name), None)
            return jsonify({**c, 'progress': progress})
    return jsonify({'error': 'not found'}), 404

@app.route('/api/client/add', methods=['POST'])
def api_client_add():
    """新增/更新客户 → 写盘 → 下次请求自动同步（出模板即入库）"""
    data = request.json or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'error': '客户名称为必填'}), 400
    template = data.get('template', '').strip()
    if not template:
        return jsonify({'error': '模板文本为必填'}), 400
    client = {
        'name': name,
        'tag': data.get('tag', ''),
        'income': data.get('income'),
        'tax': data.get('tax'),
        'biz_inst_count': data.get('biz_inst_count', 0) or 0,
        'bank_count': data.get('bank_count', 0) or 0,
        'all_institutions': data.get('all_institutions', []),
        'template': template,
        'created': data.get('created', datetime.now().isoformat()[:10]),
        'status': data.get('status', '历史'),
        # 结构化风控字段（模板数据→银行判断标准）
        'legal_age': data.get('legal_age'),
        'tax_grade': data.get('tax_grade', ''),
        'debt_total': data.get('debt_total'),
        'query_3m': data.get('query_3m'),
        'query_6m': data.get('query_6m'),
        'query_12m': data.get('query_12m'),
        'cc_usage': data.get('cc_usage'),
        'sale_ratio': data.get('sale_ratio'),
        'has_overdue': data.get('has_overdue', False),
        'overdue_cont': data.get('overdue_cont', 0) or 0,
        'overdue_cum': data.get('overdue_cum', 0) or 0,
        'established_year': data.get('established_year'),
        'shareholding_pct': data.get('shareholding_pct'),
        'is_dual_sign': data.get('is_dual_sign', False),
    }
    # 从模板文本自动解析回填空缺字段
    parsed = parse_template(template)
    for k, v in parsed.items():
        if client.get(k) in (None, '', 0, False):
            client[k] = v
    # 派生字段
    tag = client.get('tag', '')
    client['is_gaoxin'] = '国高新' in tag
    client['has_tech'] = any(t in tag for t in ['国高新', '省专精特新', '专精特新', '科技型', '创新型'])
    # 以模板为单一真相源：smart提取 income/tax/biz_inst_count 覆盖前端传入值
    # （前端分区表单填的值可能与模板不一致时，以模板为准）
    tpl_income = extract_income_from_tpl(template)
    if tpl_income is not None:
        client['income'] = tpl_income
    tpl_tax = extract_tax_from_tpl(template)
    if tpl_tax is not None:
        client['tax'] = tpl_tax
    tpl_inst = extract_biz_inst_from_tpl(template)
    if tpl_inst is not None:
        client['biz_inst_count'] = tpl_inst
    existing = next((c for c in WB['clients'] if c['name'] == name), None)
    if existing:
        existing.update(client)
        action = 'updated'
    else:
        WB['clients'].append(client)
        WB['total_clients'] = len(WB['clients'])
        action = 'added'
    # 同步在做进展
    if data.get('active') and data.get('progress'):
        prog = next((a for a in ACTIVE if a['name'] == name), None)
        if not prog:
            ACTIVE.append({
                'name': name,
                'status': data.get('status_progress', '待反馈'),
                'deadline': data.get('deadline'),
                'overdue': data.get('overdue', False),
                'progress': data.get('progress', ''),
                'next': data.get('next', ''),
            })
        else:
            prog.update({
                'status': data.get('status_progress', prog.get('status', '待反馈')),
                'deadline': data.get('deadline', prog.get('deadline')),
                'progress': data.get('progress', prog.get('progress', '')),
                'next': data.get('next', prog.get('next', '')),
            })
    save_wb()
    load_all()  # 写后立即刷新，保证返回一致
    return jsonify({'action': action, 'total_clients': WB['total_clients']})

# ============ 银行 / 产品 ============
@app.route('/api/banks')
def api_banks():
    tier = request.args.get('tier', '')
    result = BANKS
    if tier == '红线':
        result = [b for b in result if b['red']]
    elif tier:
        result = [b for b in result if b['tier'] == tier]
    out = [{k: v for k, v in b.items() if k != 'products'} for b in result]
    return jsonify(out)

@app.route('/api/bank/<path:name>')
def api_bank_detail(name):
    for b in BANKS:
        if b['name'] == name:
            return jsonify(b)
    return jsonify({'error': 'not found'}), 404

@app.route('/api/bank/add', methods=['POST'])
def api_bank_add():
    """新增/更新一款银行产品 → 写盘 → 自动同步。银行政策改了即生效，无需重启"""
    data = request.json or {}
    bank = data.get('bank', '').strip()
    product = data.get('product', '').strip()
    if not bank or not product:
        return jsonify({'error': '银行和 product 名称为必填'}), 400
    rec = {'bank': bank, 'product': product}
    for k in ('type', 'amount_max', 'rate', 'min_age', 'max_age', 'tax_grades',
              'min_tax', 'min_established', 'query_3m_max', 'query_6m_max',
              'inst_max', 'debt_max', 'overdue_rule', 'min_shareholding',
              'dual_sign_required', 'rate_min', 'rate_max'):
        if k in data:
            rec[k] = data[k]
    existing = next((p for p in PRODUCTS if p['bank'] == bank and p['product'] == product), None)
    if existing:
        existing.update(rec)
        action = 'updated'
    else:
        PRODUCTS.append(rec)
        action = 'added'
    save_products()
    load_all()
    return jsonify({'action': action, 'total_products': len(PRODUCTS)})

@app.route('/api/rules/update', methods=['POST'])
def api_rules_update():
    """更新分层/获贷率/红线规则 → 写盘 → 自动同步。改规则不用动代码"""
    data = request.json or {}
    if 'tiers' in data:
        RULES['tiers'] = data['tiers']
    if 'rates' in data:
        RULES['rates'] = data['rates']
    if 'redline' in data:
        RULES['redline'] = data['redline']
    if 'tier_order' in data:
        RULES['tier_order'] = data['tier_order']
    if 'tier_bonus' in data:
        RULES['tier_bonus'] = data['tier_bonus']
    save_rules()
    load_all()
    return jsonify({'ok': True, 'tiers': len(RULES.get('tiers', {})),
                    'rates': len(RULES.get('rates', {})), 'redline': RULES.get('redline', [])})

@app.route('/api/rules')
def api_rules():
    """读取当前规则全量（编辑界面用）"""
    return jsonify(RULES)

@app.route('/api/active/update', methods=['POST'])
def api_active_update():
    """更新在做客户进展 → 写入 workbench_data.json 的 active[] → 自动同步"""
    data = request.json or {}
    name = data.get('name', '').strip()
    if not name:
        return jsonify({'error': '客户名称为必填'}), 400
    prog = next((a for a in ACTIVE if a['name'] == name), None)
    if not prog:
        prog = {'name': name, 'status': '待反馈', 'deadline': None,
                'overdue': False, 'progress': '', 'next': ''}
        ACTIVE.append(prog)
    for k in ('status', 'deadline', 'overdue', 'progress', 'next'):
        if k in data:
            prog[k] = data[k]
    WB['active'] = ACTIVE
    save_wb()
    load_all()
    return jsonify({'ok': True, 'active_count': len(ACTIVE)})

@app.route('/api/reload', methods=['POST'])
def api_reload():
    """手动触发全量重载（如直接编辑了 JSON 文件，调一次即刻生效）"""
    load_all()
    return jsonify({'ok': True, 'clients': len(WB['clients']),
                    'banks': len(BANKS), 'products': len(PRODUCTS)})

def extract_income_from_tpl(tpl):
    """从模板提取年开票额——优先匹配完整年度（如"2023年开票1600万"），
    跳过非整年数据（如"2024年1-9月"）。若无整年数据则取最后一个匹配值。
    匹配优先级：开票XXX万 > 应税销售额XXX(万/亿) > 营收XXX万（排除"利润表营收"）。
    支持"约/近/约"等修饰词；支持经营数据跨多行（集团口径）。"""
    if not tpl:
        return None
    # 收集带年份上下文的数据和不带年份的数据
    full_year = []   # 整年数据：如"2023年开票1600万"
    partial_year = []  # 非整年数据：如"2024年1-9月开票约1000万"
    no_year = []     # 无年份的数据

    # 全局匹配（不限于"经营数据"单行，因为有的客户经营数据跨多行写集团口径）
    # 1. 开票XXX万（支持"约/近"），同时检查是否是整年
    for m in re.finditer(r'(?:(\d{4})\s*年\s*(?:1-\d{1,2}月)?)?\s*开票(?:约|近|约)?\s*([\d.]+)\s*万', tpl):
        year_hint = m.group(1)
        val = float(m.group(2))
        if val <= 0:
            continue
        # 判断是否是整年：有年份且没有"1-X月"模式
        full_match = m.group(0)
        is_partial = bool(re.search(r'(\d{4})\s*年\s*\d{1,2}-\d{1,2}月', full_match))
        if year_hint and not is_partial:
            full_year.append((int(year_hint), val))
        elif is_partial:
            partial_year.append(val)
        else:
            no_year.append(val)

    # 2. 应税销售额XXX亿
    for m in re.finditer(r'应税销售额?([\d.]+)\s*亿', tpl):
        no_year.append(round(float(m.group(1)) * 10000, 1))
    # 3. 应税销售额XXX万
    for m in re.finditer(r'应税销售额?([\d.]+)\s*万', tpl):
        val = float(m.group(1))
        if val > 0:
            no_year.append(val)
    # 4. 营收XXX万（排除"利润表营收"）支持"约"
    for m in re.finditer(r'(?<!利润表)营收(?:约|近|约)?\s*([\d.]+)\s*万', tpl):
        val = float(m.group(1))
        if val > 0:
            no_year.append(val)
    # 5. 兜底：流水/营业额（当开票/营收/应税销售额都没有时）
    if not full_year and not partial_year and not no_year:
        # 5a. 营业额一年X亿+ → 转为万
        for m in re.finditer(r'营业额(?:一年|年)?(?:约|近)?\s*([\d.]+)\s*亿', tpl):
            no_year.append(round(float(m.group(1)) * 10000, 1))
        # 5b. 营业额XXX万（支持"约"）
        for m in re.finditer(r'营业额(?:一年|年)?(?:约|近)?\s*([\d.]+)\s*万', tpl):
            no_year.append(float(m.group(1)))
        # 5c. 流水约XXX万/年 或 年流水约XXX万
        if not no_year:
            for m in re.finditer(r'(?:年\s*)?流水(?:约|近|约)?\s*([\d.]+)\s*万', tpl):
                val = float(m.group(1))
                if val > 0:
                    no_year.append(val)
        # 5d. 年含个人卡流水约XXX万
        if not no_year:
            for m in re.finditer(r'年含个人卡流水(?:约|近|约)?\s*([\d.]+)\s*万', tpl):
                no_year.append(float(m.group(1)))
        # 5e. 个人微信年收入约XX万（个人客户兜底）
        if not no_year:
            for m in re.finditer(r'年收入(?:约|近|约)?\s*([\d.]+)\s*万', tpl):
                no_year.append(float(m.group(1)))

    # 优先级：整年（取最新年份） > 非整年（取最后一个） > 无年份（取最后一个）
    if full_year:
        full_year.sort(key=lambda x: x[0])  # 按年份排序
        return full_year[-1][1]  # 取最新年份的值
    if partial_year:
        return partial_year[-1]
    if no_year:
        return no_year[-1]
    return None


def extract_tax_from_tpl(tpl):
    """从模板提取年纳税——取最新年度"""
    biz_section = ''
    for line in tpl.split('\n'):
        if '经营数据' in line:
            biz_section = line
            break
    if not biz_section:
        biz_section = tpl
    taxes = []
    for m in re.finditer(r'纳税(?:约|近)?\s*([\d.]+)\s*万', biz_section):
        val = float(m.group(1))
        if val > 0:
            taxes.append(val)
    if taxes:
        return taxes[-1]
    return None


def extract_biz_inst_from_tpl(tpl):
    """从模板提取企业端机构数——只从"机构数"行提取，支持"约/近"前缀"""
    for line in tpl.split('\n'):
        stripped = line.strip()
        if stripped.startswith('机构数'):
            m = re.search(r'企业端(?:约|近|大概)?\s*(\d+)\s*家', stripped)
            if m:
                return int(m.group(1))
    return None


# 银行简称→标准名映射
BANK_ALIASES = {
    '工行': '工商银行', '建行': '建设银行', '农行': '农业银行', '中行': '中国银行',
    '交行': '交通银行', '邮储': '邮储银行', '招行': '招商银行', '中信': '中信银行',
    '浦发': '浦发银行', '民生': '民生银行', '兴业': '兴业银行', '光大': '光大银行',
    '华夏': '华夏银行', '平安': '平安银行', '浙商': '浙商银行', '渤海': '渤海银行',
    '广发': '广发银行', '恒丰': '恒丰银行',
    '新网': '新网银行', '微众': '微众银行', '网商': '网商银行',
    '宁波': '宁波银行', '杭州': '杭州银行', '南京': '南京银行',
    '江苏': '江苏银行', '北京': '北京银行', '上海': '上海银行',
    '杭州联合': '杭州联合农商行', '余杭农商': '余杭农商行', '萧山农商': '萧山农商行',
    '临安农商': '临安农商行', '泰隆': '泰隆银行', '民泰': '民泰银行',
    '温州': '温州银行', '稠州': '稠州银行',
}


def extract_institutions_from_tpl(tpl):
    """从模板提取在贷银行名单，返回去重后的list。
    只取"企业负债"段的银行（不包含"个人负债"段），因为：
    - 企业负债=对公贷款机构（银行授信）
    - 个人负债=个人消费贷/抵押贷（非对公机构）

    匹配策略：先用已知银行全名匹配，再用简称映射匹配。"""
    if not tpl:
        return []
    # 只收集"企业负债"到下一个段落标题（不包括"个人负债"）
    debt_section = []
    capturing = False
    for line in tpl.split('\n'):
        stripped = line.strip()
        if '企业负债' in stripped:
            capturing = True
            debt_section.append(stripped)
            continue
        if capturing:
            # 遇到"个人负债"或下一个段落标题则停止
            if re.match(r'^(个人负债|企业担保|机构数|法诉|征信|落实建议|资产：|经营数据|公司架构|企业标签|客户画像|公司名称)', stripped):
                break
            debt_section.append(stripped)
    debt_text = '\n'.join(debt_section)

    found = set()
    # 1. 先匹配已知银行全名（从WB的insts列表）
    for inst in WB.get('insts', []):
        bname = inst['name']
        if bname in debt_text:
            found.add(bname)
    # 2. 匹配简称→全名
    for alias, full in BANK_ALIASES.items():
        if alias in debt_text:
            found.add(full)
    # 3. 匹配"XX银行"模式（兜底，排除"银行承兑汇票"等非贷款机构）
    for m in re.finditer(r'([\u4e00-\u9fa5]{2,6})银行', debt_text):
        name = m.group(1) + '银行'
        found.add(name)
    # 4. 匹配"XX农商"模式
    for m in re.finditer(r'([\u4e00-\u9fa5]{2,4})农商', debt_text):
        name = m.group(1) + '农商行'
        found.add(name)

    # 过滤掉明显不是贷款机构的误匹配
    noise = {'银行承兑汇票', '银行机构数', '银行审批', '银行偏好', '银行看',
             '股份制银行', '国有大行', '本土城商行', '城商行', '大行',
             '农商银行', '商业银行', '银行'}
    found = sorted(f for f in found if f not in noise and len(f) >= 4)

    # 合并"XX农商行"和"XX农商银行"重复（保留简称版）
    deduped = []
    for f in found:
        if f.endswith('农商银行') and f.replace('农商银行', '农商行') in found:
            continue  # 跳过长名版本，保留短名版本
        if f.endswith('联合农商银行') and f.replace('联合农商银行', '联合农商行') in found:
            continue
        deduped.append(f)

    return deduped


def sync_client_from_template(c):
    """以模板文本为单一真相源，全量同步客户结构化字段（覆盖模式）。
    每次调用都从模板重新解析，确保结构化字段与模板零偏差。"""
    tpl = c.get('template', '') or ''
    if not tpl:
        return False
    changed = False

    # parse_template 提取的15个字段：覆盖
    parsed = parse_template(tpl)
    for k, v in parsed.items():
        if c.get(k) != v:
            c[k] = v
            changed = True

    # income / tax / biz_inst_count：smart提取覆盖
    tpl_income = extract_income_from_tpl(tpl)
    if tpl_income is not None:
        old = c.get('income')
        if old is None or abs((old or 0) - tpl_income) > 1:
            c['income'] = tpl_income
            changed = True
    tpl_tax = extract_tax_from_tpl(tpl)
    if tpl_tax is not None:
        old = c.get('tax')
        if old is None or abs((old or 0) - tpl_tax) > 1:
            c['tax'] = tpl_tax
            changed = True
    tpl_inst = extract_biz_inst_from_tpl(tpl)
    if tpl_inst is not None:
        if c.get('biz_inst_count') != tpl_inst:
            c['biz_inst_count'] = tpl_inst
            changed = True

    # 派生字段
    tag = c.get('tag', '') or ''
    is_gaoxin = '国高新' in tag
    has_tech = any(t in tag for t in ['国高新', '省专精特新', '专精特新', '科技型', '创新型'])
    if c.get('is_gaoxin') != is_gaoxin:
        c['is_gaoxin'] = is_gaoxin
        changed = True
    if c.get('has_tech') != has_tech:
        c['has_tech'] = has_tech
        changed = True

    # 在贷银行名单：从模板"企业负债"+"个人负债"段落提取（以模板为准，每次重新提取覆盖）
    tpl_insts = extract_institutions_from_tpl(tpl)
    if tpl_insts:
        if c.get('all_institutions') != tpl_insts:
            c['all_institutions'] = tpl_insts
            c['bank_count'] = len(tpl_insts)
            changed = True
    elif not c.get('all_institutions'):
        # 模板提取不到且已有也为空，保持空
        pass

    return changed


@app.route('/api/backfill', methods=['POST'])
def api_backfill():
    """全量同步：以模板为单一真相源，重新解析所有客户的结构化字段（覆盖模式）。
    模板更新后调用此接口，income/tax/biz_inst_count等全部与模板对齐。"""
    filled = 0
    for c in WB['clients']:
        if sync_client_from_template(c):
            filled += 1
    save_wb()
    load_all()
    return jsonify({'ok': True, 'total': len(WB['clients']), 'synced': filled})

@app.route('/api/tags')
def api_tags():
    tags = set()
    for c in WB['clients']:
        for t in (c.get('tag') or '').split('、'):
            t = t.strip()
            if t:
                tags.add(t)
    return jsonify(sorted(tags))

@app.route('/api/all_banks')
def api_all_banks():
    banks_list = sorted(set(c['name'] for c in WB.get('insts', [])))
    return jsonify(banks_list)

# ============ 智能匹配引擎 ============
@app.route('/api/match', methods=['POST'])
def api_match():
    """输入客户条件，返回可推银行排序 + 匹配理由 + 排除理由"""
    return jsonify(run_match(request.json or {}))

def run_match(data):
    age = data.get('age')
    income = data.get('income')
    tax = data.get('tax')
    tax_grade = data.get('tax_grade', '')
    query_3m = data.get('query_3m')
    query_6m = data.get('query_6m')
    query_12m = data.get('query_12m')
    inst_count = data.get('inst_count')
    debt_total = data.get('debt_total')
    cc_usage = data.get('cc_usage')
    overdue = data.get('overdue', False)
    overdue_cont = data.get('overdue_cont')
    overdue_cum = data.get('overdue_cum')
    sale_ratio = data.get('sale_ratio')
    has_tech = data.get('has_tech', False)
    is_gaoxin = data.get('is_gaoxin', False)
    exclude_banks = set(data.get('exclude_banks', []))

    risk_fatal = []
    risk_warn = []
    # income 参与销贷比自动计算：如未提供 sale_ratio 但有 income 和 debt_total
    if not sale_ratio and income and debt_total and income > 0:
        sale_ratio = round(debt_total / income * 100, 1)
    if overdue_cont and overdue_cont >= 3:
        risk_fatal.append(f'连三：近2年连续逾期 {overdue_cont} 次（≥3 致命红线）')
    if overdue_cum and overdue_cum >= 6:
        risk_fatal.append(f'累六：近2年累计逾期 {overdue_cum} 次（≥6 致命红线）')
    if cc_usage and cc_usage > 90:
        risk_warn.append(f'信用卡使用率 {cc_usage}% 远超90%红线，几乎无法新增信用贷')
    elif cc_usage and cc_usage > 70:
        risk_warn.append(f'信用卡使用率 {cc_usage}% 超70%高风险线，部分银行直接拒')
    if sale_ratio and sale_ratio > 50:
        risk_warn.append(f'销贷比 {sale_ratio}% 超50%（工行等硬性指标，额度受限）')
    if query_6m and query_6m > 12:
        risk_warn.append(f'近6月查询 {query_6m} 次超12次（工行红线，线下产品可沟通）')
    if query_12m and query_12m > 20:
        risk_warn.append(f'近12月查询 {query_12m} 次超20次（多数银行重点关注）')
    if inst_count is not None and inst_count >= 3:
        risk_warn.append(f'个人信用类经营贷机构数 {inst_count} 家（工行硬限≤3，需整合）')

    tier_bonus = RULES.get('tier_bonus', {'T1': 15, 'T2': 8, 'T3': 3, 'T4': 1})
    results = []

    for b in BANKS:
        bank_name = b['name']
        reasons_match = []
        reasons_exclude = []
        score = 0
        if b['red']:
            reasons_exclude.append('红线禁推银行')
        if bank_name in exclude_banks:
            reasons_exclude.append('在排除清单中（存量/征信查询/已还清）')
        for p in b['products']:
            p_match = []
            p_exclude = []
            if age:
                if p.get('max_age') and age > p['max_age']:
                    p_exclude.append(f'年龄{age}超上限{p["max_age"]}')
                elif p.get('min_age') and age < p['min_age']:
                    p_exclude.append(f'年龄{age}低于下限{p["min_age"]}')
                else:
                    p_match.append('年龄符合')
                    score += 5
            if tax and p.get('min_tax'):
                if tax < p['min_tax']:
                    p_exclude.append(f'纳税{tax}万低于要求{p["min_tax"]}万')
                else:
                    p_match.append(f'纳税{tax}万达标')
                    score += 10
            if tax_grade and p.get('tax_grades'):
                if tax_grade in p['tax_grades']:
                    p_match.append(f'纳税{tax_grade}级符合')
                    score += 5
            if query_3m and p.get('query_3m_max'):
                if query_3m > p['query_3m_max']:
                    p_exclude.append(f'近3月查询{query_3m}>上限{p["query_3m_max"]}')
                else:
                    score += 8
            if query_6m and p.get('query_6m_max'):
                if query_6m > p['query_6m_max']:
                    p_exclude.append(f'近6月查询{query_6m}>上限{p["query_6m_max"]}')
                else:
                    score += 8
            if inst_count is not None and p.get('inst_max'):
                if inst_count > p['inst_max']:
                    p_exclude.append(f'机构数{inst_count}>上限{p["inst_max"]}')
                else:
                    score += 8
            if debt_total and p.get('debt_max'):
                if debt_total > p['debt_max']:
                    p_exclude.append(f'总负债{debt_total}万>上限{p["debt_max"]}万')
                else:
                    score += 5
            if p.get('tech_required') and not (has_tech or is_gaoxin):
                p_exclude.append('需科技标签但客户无')
            elif p.get('tech_required') and (has_tech or is_gaoxin):
                p_match.append('科技标签符合')
                score += 15
            if overdue:
                p_exclude.append('有逾期记录')
                score -= 20
            if cc_usage and cc_usage > 50 and '工商银行' in bank_name:
                p_exclude.append(f'信用卡使用率{cc_usage}%>工行50%限制')
            rate_str = b.get('rate', '')
            if rate_str:
                try:
                    rate_val = float(rate_str.replace('%', ''))
                    score += rate_val * 0.5
                except Exception:
                    pass
            score += tier_bonus.get(b['tier'], 0)
            if p_match and not p_exclude:
                reasons_match.append(f'{p["product"]}: {"、".join(p_match)}')
            elif p_exclude:
                reasons_exclude.append(f'{p["product"]}: {"、".join(p_exclude)}')
        if risk_fatal:
            continue
        if reasons_match and not [e for e in reasons_exclude if '红线' in e or '排除清单' in e]:
            # 提取匹配产品名（match_reasons格式为"产品名: 理由"，取第一个匹配产品的名称）
            best_product = ''
            if reasons_match:
                best_product = reasons_match[0].split(':')[0].strip()
            results.append({
                'bank': bank_name,
                'tier': b['tier'],
                'rate': b['rate'],
                'score': round(score, 1),
                'match_reasons': reasons_match[:5],
                'exclude_reasons': [e for e in reasons_exclude if '红线' not in e and '排除清单' not in e][:3],
                'product_count': b['product_count'],
                'best_product': best_product,
            })
    results.sort(key=lambda x: x['score'], reverse=True)
    return {
        'total_match': len(results),
        'recommendations': results[:5],
        'all_matches': results,
        'risk_fatal': risk_fatal,
        'risk_warn': risk_warn,
    }

# ============ 诊断报告导出 ============
def _extract_advice(template_text):
    """从模板文本中提取"落实建议"段落（四段式），返回list[str]。"""
    if not template_text:
        return []
    lines = template_text.split('\n')
    advice_lines = []
    capturing = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if '落实建议' in stripped or '建议方案' in stripped or '操作建议' in stripped:
            capturing = True
            # 如果标题行后面还有内容，也收集
            after_colon = stripped.split('：', 1)[-1].split(':', 1)[-1].strip()
            if after_colon and after_colon != stripped:
                advice_lines.append(after_colon)
            continue
        if capturing:
            # 遇到下一个大段落标题则停止
            if re.match(r'^[一二三四五六七八九十]、', stripped) or re.match(r'^第[一二三四五六七八九十]段', stripped):
                break
            advice_lines.append(stripped)
    return advice_lines if advice_lines else []


@app.route('/api/report/<path:name>')
def api_report(name):
    """导出客户融资诊断报告 DOCX（六板块：画像速览/风控预检/匹配Top5/排除清单/落实建议/完整档案）"""
    c = None
    for cl in WB['clients']:
        if cl['name'] == name:
            c = cl
            break
    if not c:
        return jsonify({'error': 'not found'}), 404
    tag = c.get('tag', '') or ''

    # ---- 构造匹配引擎输入（17字段全量 + parse_template回退）----
    # 优先用客户已有结构化字段，空缺则从模板文本实时解析补全
    parsed = parse_template(c.get('template', ''))
    def _g(key, default=None):
        """先查客户字段，空则查parse_template回退，再空则default"""
        v = c.get(key)
        if v is None or v == '':
            v = parsed.get(key)
        return v if v is not None else default

    body = {
        'age': _g('legal_age'),
        'income': c.get('income'),
        'tax': c.get('tax'),
        'tax_grade': _g('tax_grade', ''),
        'query_3m': _g('query_3m'),
        'query_6m': _g('query_6m'),
        'query_12m': _g('query_12m'),
        'inst_count': c.get('biz_inst_count'),
        'debt_total': _g('debt_total'),
        'cc_usage': _g('cc_usage'),
        'sale_ratio': _g('sale_ratio'),
        'overdue': _g('has_overdue', False),
        'overdue_cont': _g('overdue_cont'),
        'overdue_cum': _g('overdue_cum'),
        'exclude_banks': c.get('all_institutions', []),
        'is_gaoxin': '国高新' in tag,
        'has_tech': any(t in tag for t in ['国高新', '省专精特新', '科技型', '创新型']),
    }
    match_res = run_match(body)

    # ---- 生成 DOCX 报告 ----
    doc = Document()
    doc.add_heading('融资诊断报告', 0)
    p = doc.add_paragraph()
    p.add_run(c['name']).bold = True
    doc.add_paragraph('生成时间：' + datetime.now().isoformat()[:19])
    doc.add_heading('一、客户画像速览', level=1)
    info = doc.add_table(rows=1, cols=4)
    info.style = 'Light Grid Accent 1'
    cells = info.rows[0].cells
    cells[0].text = '客户名称'
    cells[1].text = c['name']
    cells[2].text = '企业标签'
    cells[3].text = tag or '无'

    def _fmt(v, suffix=''):
        if v is None or v == '':
            return '待补充'
        return f'{v}{suffix}'

    rows_data = [
        ('年开票(万)', _fmt(c.get('income')), '年纳税(万)', _fmt(c.get('tax'))),
        ('法人年龄', _fmt(_g('legal_age'), '岁'), '纳税级别', _fmt(_g('tax_grade'), '级')),
        ('企业总负债(万)', _fmt(_g('debt_total')), '信用卡使用率', _fmt(_g('cc_usage'), '%')),
        ('销贷比', _fmt(_g('sale_ratio'), '%'), '近12月查询', _fmt(_g('query_12m'), '次')),
        ('近3月查询', _fmt(_g('query_3m'), '次'), '近6月查询', _fmt(_g('query_6m'), '次')),
        ('逾期情况', '有逾期' if _g('has_overdue') else '无逾期', '连三', '是' if (_g('overdue_cont') or 0) >= 3 else '否'),
        ('累六', '是' if (_g('overdue_cum') or 0) >= 6 else '否', '在贷银行数', _fmt(c.get('bank_count'), '家')),
        ('个人信用类机构数', _fmt(c.get('biz_inst_count'), '家'), '企业标签', tag or '无'),
    ]
    for row_vals in rows_data:
        r = info.add_row().cells
        r[0].text = row_vals[0]
        r[1].text = row_vals[1]
        r[2].text = row_vals[2]
        r[3].text = row_vals[3]
    # ---- 板块二：风控预检结论 ----
    has_risk = bool(match_res.get('risk_fatal') or match_res.get('risk_warn'))
    doc.add_heading('二、风控预检结论', level=1)
    if match_res.get('risk_fatal'):
        p = doc.add_paragraph()
        p.add_run('致命红线（一票否决）：').bold = True
        for r in match_res['risk_fatal']:
            doc.add_paragraph(r, style='List Bullet')
    if match_res.get('risk_warn'):
        p = doc.add_paragraph()
        p.add_run('高风险项：').bold = True
        for r in match_res['risk_warn']:
            doc.add_paragraph(r, style='List Bullet')
    if not has_risk:
        doc.add_paragraph('未触发致命红线和高风险项，风控预检通过。')

    # ---- 板块三：智能匹配Top5推荐银行 ----
    recs = match_res.get('recommendations', [])
    if recs and not match_res.get('risk_fatal'):
        doc.add_heading('三、智能匹配Top5推荐银行', level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = '排序'
        hdr[1].text = '银行'
        hdr[2].text = '分层'
        hdr[3].text = '推荐产品'
        hdr[4].text = '匹配理由摘要'
        for i, r in enumerate(recs, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = r.get('bank', '')
            row[2].text = r.get('tier', '')
            row[3].text = r.get('best_product', '')
            row[4].text = '；'.join(r.get('match_reasons', [])[:3])
    elif match_res.get('risk_fatal'):
        doc.add_heading('三、建议处理方向', level=1)
        doc.add_paragraph('命中致命红线，建议先处理逾期记录或转向抵押/担保路径，再重新评估。')

    # ---- 板块四：排除清单 ----
    doc.add_heading('四、排除清单（存量/征信/已还清银行）', level=1)
    if c.get('all_institutions'):
        doc.add_paragraph('、'.join(c['all_institutions']))
    else:
        doc.add_paragraph('无排除银行。')

    # ---- 板块五：落实建议（四段式，从模板提取）----
    doc.add_heading('五、落实建议', level=1)
    advice_list = _extract_advice(c.get('template', ''))
    if advice_list:
        for i, advice in enumerate(advice_list, 1):
            doc.add_paragraph(advice, style='List Number')
    else:
        doc.add_paragraph('模板中暂无落实建议段落，请顾问根据风控结论和匹配结果手动补充。')

    # ---- 板块六：完整客户档案原文 ----
    doc.add_heading('六、完整客户档案', level=1)
    doc.add_paragraph(c.get('template', '（暂无模板文本）'))
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name=c['name'] + '_融资诊断报告.docx')

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8000))
    app.run(host='0.0.0.0', port=port, debug=False)
